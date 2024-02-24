import pandas as pd
from docx import Document
import docx 
import psycopg2
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml import ns
from openpyxl.utils.dataframe import dataframe_to_rows
from docx.oxml.ns import qn 
from openpyxl.styles import Font
import openpyxl
import xlwings as xw
import subprocess
import os
import sys
db_params = {
    'dbname': 'secdata_genai_newdevhub_v3',
    'user': 'postgres',
    'password': 'FcstNew#2023',
    'host': 'aiml-dev-db.cbtaanq5yhb4.us-east-1.rds.amazonaws.com',
    'port': '5432'
}
#parameter1 = sys.argv[1]
parameter1 ='def'
# Function to update Word document under a specific heading
def add_hyperlink(paragraph, text, url):
    # Add a hyperlink to a paragraph
    run = paragraph.add_run(text)

    if url is not None:
        # Create a hyperlink element
        hyperlink = OxmlElement('w:hyperlink')

        # Add a relationship for the hyperlink
        rel_id = paragraph.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink.set(qn('r:id'), rel_id)

        # Add the hyperlink element to the paragraph
        run._r.insert(0, hyperlink)

        # Set additional hyperlink properties
        run.font.color.rgb = RGBColor(0, 0, 255)  # Set text color to blue
        run.font.underline = True
        
# def get_hyperlink_from_excel(file_path, sheet_name, row_num, col_num):
#     workbook = openpyxl.load_workbook(file_path, data_only=True)
#     sheet = workbook[sheet_name]

#     cell = sheet.cell(row=row_num, column=col_num)

#     if hasattr(cell, 'hyperlink') and cell.hyperlink:
#         # If the cell has a hyperlink, return the hyperlink address
#         return cell.hyperlink.target

#     return None

def get_hyperlink_from_excel(file_path, sheet_name, row_num, col_num):
    # Open the workbook
    workbook = xw.Book(file_path)

    # Open the sheet
    sheet = workbook.sheets[sheet_name]

    # Get the cell
    cell = sheet.cells[row_num, col_num]

    # Check if the cell has a hyperlink
    if cell.hyperlink:
        return cell.hyperlink.address

    return None

def update_word_document(word_file_path, heading, data,initiativename,account):
    if os.path.exists(word_file_path):
        # If the file exists, append content to it
        doc = Document(word_file_path)
    else:
        # If the file doesn't exist, create a new document
        doc = Document()
    #print(data.columns)
    print(heading)
    header_found = False

    #print('debug1')
    initiativename_only=initiativename[initiativename.index('.')+2:initiativename.index('.')+256]
    str_content_1 = initiativename + ':' 
    str_content_2 = 'Below table shows summary of documents to refer when discussing about initiative for '+initiativename_only+':'
            #print(str_content)
            ##paragraph.insert_paragraph_before(str_content)
    #doc.add_paragraph(str_content)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('\n'+str_content_1)
    run.bold = True 
    if data.empty:
        str_content_2 = 'Please add the Use Cases/Collaterals/Case Study to map with this Initiative'
        doc.add_paragraph(str_content_2)
    else:
        doc.add_paragraph(str_content_2)
        doc.add_table(data.shape[0] + 1, data.shape[1]).style = 'Table Grid'
        table = doc.tables[-1]
        column_widths_inches = [0.25, 0.25, 0.25,5.0,0.5,5.0,10.0]  # Widths in inches
        column_widths_twentieths = [int(width * 72 * 20) for width in column_widths_inches]

        for i, width in enumerate(column_widths_twentieths):
            table.columns[i].width = width
    #run = paragraph.add_run('\n')

            # Write column headers
        for col_num, col_name in enumerate(data.columns):
            table.cell(0, col_num).text = col_name
            cell = table.cell(0, col_num) 
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            font = run.font
            font.size = Pt(8)
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 255)  # Set text color to white
        
                #print(col_name)

            # Write data to the table
            for row_num in range(data.shape[0]):
                hyperlink_added = False
                for col_num, value in enumerate(data.iloc[row_num]):
                    if data.columns[col_num] == 'document':
                        hyperlink_address = get_hyperlink_from_excel(excel_file_path, sheet_name, row_num+2, col_num+3)
                        cell = table.cell(row_num + 1, col_num)
                        paragraph = cell.paragraphs[0]
                        add_hyperlink(paragraph, str(value), hyperlink_address)
                    else:
                        table.cell(row_num + 1, col_num).text = str(value)
                        cell = table.cell(row_num + 1, col_num)
                        paragraph = cell.paragraphs[0]
                        run = paragraph.runs[0]
                        font = run.font
                        font.size = Pt(8)
 
    #run = paragraph.add_run('\n')
    # Save the modified Word document
    #doc.save(word_file_path.replace('.docx', '_updated.docx'))
    print(word_file_path)
    doc.save(word_file_path)

query = f"SELECT distinct doc_path FROM dwh.f_initiative_bkup where accountid ='{parameter1}'"
print(query)
connection = psycopg2.connect(**db_params)
cursor = connection.cursor()
cursor.execute(query)
data = cursor.fetchall()
print(data)
word_file_path = data[0]
word_file_path_1="".join(word_file_path)
#temp_file_path= word_file_path_1.replace('.docx','_temp.docx')
#temp_file_path_1="".join(temp_file_path)

# Replace 'your_excel_file.xlsx' with the path to your Excel file
excel_file_path = 'C:/Users/Pranjal/Desktop/PRANJAL/Main/Pranjal/gen_ai/SecData_GenAI_NewDevHub_version3/DS/ai_framework/NextQSummary/data/df_mappings_20240206131040_Overall.xlsx'

# Replace 'your_word_file.docx' with the path to your Word file
word_file_path = 'C:/Users/Pranjal/Desktop/PRANJAL/Main/Pranjal/gen_ai/SecData_GenAI_NewDevHub_version3/DS/ai_framework/NextQSummary/data/lakshmi/Icelandair-Insights_through_GenAI.docx'
temp_file_path = 'C:/Users/Pranjal/Desktop/PRANJAL/Main/Pranjal/gen_ai/SecData_GenAI_NewDevHub_version3/DS/ai_framework/NextQSummary/data/lakshmi/Icelandair-Insights_through_GenAI_temp.docx'

# Replace 'your_sheet_name' with the name of your sheet
sheet_name = 'Sheet1'

# Read the Excel file into a pandas DataFrame
#df = pd.read_excel(excel_file_path, sheet_name=sheet_name,usecols='F:L')
df = pd.read_excel(excel_file_path,
                   sheet_name='Sheet1',
                   usecols='C,D,F:L',
                   header=0,
                   engine='openpyxl') 
connection = psycopg2.connect(**db_params)
cursor = connection.cursor()
accountname ='Icelandair'
query1 = f"SELECT distinct accountname from dwh.f_initiative_bkup where accountid ='{parameter1}'"
cursor.execute(query1)
    #data1 = cursor.fetchall()
returned_list = [item[0] for item in cursor.fetchall()]
    # Construct the query to retrieve specific column values
header_found = False
# print(returned_list)
for k in returned_list:
    #print(k)
    query = f"SELECT  distinct initiativename FROM dwh.f_initiative_bkup where accountname ='{k}' order by initiativename"
    cursor.execute(query)
    data = cursor.fetchall()
    #print(data)
    for index,value in enumerate(data) :
            #print(index)
            #print(value)
            if index == 0 :
                heading = 'Appendix:'
            else :
                heading = ''
            data_rw1 = list(data[index])
            initiativenum = str(index +1)
            initiativename = data_rw1[0]
            #file1 = data_rw1[1]
            print(initiativename)
#row_numbers = [x+1 for x in df[df['accountname'] == 'JPMorgan Chase'].index]
            conditions = {'accountname' : k }
            conditions_1 = {'initiativename' : initiativename }
            condition_column2 = 'initiativename'
            condition_value2 = 'desired_value2'
            df['accountname'] = df['accountname'].astype(str)
            df['initiativename'] = df['initiativename'].astype(str)
#print(row_numbers)
#df = pd.read_excel(excel_file_path, sheet_name='Sheet1',usecols='C,F:L',header=0,**conditions)
            selected_rows_df = df[(df['accountname'] == conditions['accountname']) & (df['initiativename'] == conditions_1['initiativename'])]
#df = pd.read_excel(excel_file_path, engine='openpyxl', query=conditions)
            selected_rows_cols_df =selected_rows_df.loc[:,'document_type':'solution']
            print(selected_rows_cols_df)
            initiativename_str = initiativenum + '. '+ initiativename
# Replace 'Your Heading' with the heading where you want to insert the Excel data
            update_word_document(temp_file_path, heading, selected_rows_cols_df,initiativename_str,k)
py_file_path='D:/gen_ai/SecData_GenAI_NewDevHub_version3/DS/ai_framework/secWebScrapping/scripts/test.py'
py_cmd = f"python {py_file_path} {temp_file_path} {word_file_path_1}"
print(py_cmd)
print(temp_file_path)
print(word_file_path_1)
        #p = subprocess.Popen(dwh_cmd, stdout=subprocess.PIPE, shell=True)
#subprocess.run(py_cmd.split())
