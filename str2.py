from docx import Document

def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

if __name__ == "__main__":
    # Replace 'example.docx' with the path to your Word document
    document_path = 'C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx'

    # Load the Word document
    doc = Document(document_path)

    # Specify the text to replace and the new text
    old_text = '{str_company}'
    new_text = 'JP Morgan Chase'

    # Call the replace_text function
    replace_text(doc, old_text, new_text)

    # Save the modified document
    doc.save('C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx')

# import aspose.words as aw

# # load Word document
# doc = aw.Document("C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx")

# # replace text
# doc.range.replace("{str_company}", "JP Morgan Chase", aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))

# # save the modified document
# doc.save("C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx")

# from docx import Document
# #open the document
# doc=Document('C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx')
# Dictionary = {"{str_company}": "JP Morgan Chase"}
# for i in Dictionary:
#     for p in doc.paragraphs:
#         if p.text.find(i)>=0:
#             p.text=p.text.replace(i,Dictionary[i])
# #save changed document
# doc.save('C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx')

# from docx import Document

# def replace_str_company(doc, keyword):
#     for paragraph in doc.paragraphs:
#         # Replace in paragraphs outside the table
#         if '{str_company}' in paragraph.text:
#             paragraph.text = paragraph.text.replace('{str_company}', keyword)

#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     # Replace in each paragraph inside the cell
#                     if '{str_company}' in paragraph.text:
#                         paragraph.clear()
#                         paragraph.add_run(paragraph.text.replace('{str_company}', keyword))
# # Example usage:
# doc_path = "C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx"
# keyword = "JP Morgan Chase"

# doc = Document(doc_path)
# replace_str_company(doc, keyword)

# # Save the modified document
# doc.save("C:/Users/Pranjal/Downloads/CustomerName_AccountPlanningTemplate_Cockroach Labs 1.docx")
